from __future__ import annotations

import re
import shutil
import tempfile
import zipfile
from pathlib import Path
from posixpath import normpath
from xml.etree import ElementTree as ET

from PIL import Image
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
COURSE = HERE.parents[1]
TEMPLATE = COURSE / "templates" / "尚硅谷嵌入式项目之boomPI_v4.docx"
SOURCE = HERE / "chapter04.md"
OUTPUT = HERE / "第04章_Linux系统简介_审阅稿_v0.1.docx"

BODY = "文档正文样式"
H1 = "一级标题"
H2 = "二级标题"
CAPTION = "图片样式"
FONT_CN = "宋体"
FONT_HEADING = "黑体"
FONT_WEST = "Times New Roman"
FONT_CODE = "Consolas"
ACCENT = "2F6FED"
ACCENT_DARK = "20304A"
GRAY = "61718B"


def set_font(run, *, cn=FONT_CN, west=FONT_WEST, size=10.5, bold=None, color=None):
    run.font.name = west
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), west)
    fonts.set(qn("w:hAnsi"), west)
    fonts.set(qn("w:eastAsia"), cn)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def clear_document_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure_styles(doc: Document) -> None:
    normal = doc.styles[BODY]
    normal.font.name = FONT_WEST
    normal.font.size = Pt(10.5)
    fonts = normal.element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), FONT_WEST)
    fonts.set(qn("w:hAnsi"), FONT_WEST)
    fonts.set(qn("w:eastAsia"), FONT_CN)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Pt(21)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size in ((H1, 15), (H2, 14)):
        style = doc.styles[name]
        style.font.name = FONT_WEST
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(ACCENT_DARK)
        fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
        fonts.set(qn("w:ascii"), FONT_WEST)
        fonts.set(qn("w:hAnsi"), FONT_WEST)
        fonts.set(qn("w:eastAsia"), FONT_HEADING)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12 if name == H1 else 9)
        style.paragraph_format.space_after = Pt(5)

    caption = doc.styles[CAPTION]
    caption.font.name = FONT_WEST
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(GRAY)
    fonts = caption.element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), FONT_WEST)
    fonts.set(qn("w:hAnsi"), FONT_WEST)
    fonts.set(qn("w:eastAsia"), FONT_CN)
    caption.paragraph_format.first_line_indent = None
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)


def disable_paragraph_numbering(paragraph) -> None:
    """Keep the explicit chapter/section numbers in Markdown without template numbering."""
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:numPr"))
    if existing is not None:
        p_pr.remove(existing)
    num_pr = OxmlElement("w:numPr")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "0")
    num_pr.append(num_id)
    p_pr.append(num_pr)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, size=9, color=GRAY)


def configure_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    for child in list(header._element):
        header._element.remove(child)
    hp = header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_after = Pt(2)
    run = hp.add_run("尚硅谷嵌入式项目之 boomPI（RV1106）")
    set_font(run, cn=FONT_HEADING, size=9.5, bold=True, color=ACCENT_DARK)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "B8C4D6")
    border.append(bottom)
    hp._p.get_or_add_pPr().append(border)

    footer = section.footer
    for child in list(footer._element):
        footer._element.remove(child)
    fp = footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    run = fp.add_run("第4章 Linux系统简介 · 独立审阅稿  |  ")
    set_font(run, size=9, color=GRAY)
    add_page_field(fp)


def add_inline_text(paragraph, text: str, size=10.5):
    token_re = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\[[^\]]+\]\(https?://[^)]+\)|https?://\S+)")
    cursor = 0
    for match in token_re.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_font(run, size=size)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, cn=FONT_CODE, west=FONT_CODE, size=max(8.5, size - 1), color="B5482D")
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, size=size, bold=True, color=ACCENT_DARK)
        elif token.startswith("["):
            label, url = re.fullmatch(r"\[([^\]]+)\]\((https?://[^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        else:
            add_hyperlink(paragraph, token, token)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_font(run, size=size)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), ACCENT)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT_WEST)
    fonts.set(qn("w:hAnsi"), FONT_WEST)
    fonts.set(qn("w:eastAsia"), FONT_CN)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "19")
    rpr.extend([fonts, color, underline, size])
    run.append(rpr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_callout(doc: Document, text: str):
    p = doc.add_paragraph(style=BODY)
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.left_indent = Cm(0.45)
    p.paragraph_format.right_indent = Cm(0.2)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(7)
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EAF1FF")
    ppr.append(shd)
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), ACCENT)
    pbdr.append(left)
    ppr.append(pbdr)
    add_inline_text(p, text)


def add_figure(doc: Document, image_path: str, caption_text: str):
    path = Path(image_path)
    if not path.is_absolute():
        path = HERE / path
    if not path.exists():
        raise FileNotFoundError(path)
    with Image.open(path) as image:
        width_px, height_px = image.size
    if path.name == "rv1106-block-diagram-rev18-p14.png":
        max_w = 4.6
    elif path.name == "rv1106-hardware-guide-v11-page11-rv1106-crop.png":
        max_w = 5.0
    elif path.name in {"fig-1-07a-usb-selector.png", "fig-1-07b-usb-mux.png"}:
        max_w = 4.8
    elif path.name == "fig-1-11-display-touch-fpc.png":
        max_w = 4.9
    else:
        if path.name == "os-stack.png":
            max_w = 4.8
        elif path.name == "boot-chain.png":
            max_w = 5.2
        else:
            max_w = 5.65
    max_h = 7.0
    ratio = width_px / height_px
    width = min(max_w, max_h * ratio)
    if width_px < 450:
        width = min(width, 2.3)
    p = doc.add_paragraph(style=CAPTION)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run()
    picture = run.add_picture(str(path), width=Inches(width))
    if path.name == "terminal-board-identity.png":
        # 只在DOCX显示身份区：保留 y=0..335，裁掉网络、判定句和采集提示。
        picture.height = Inches(width * 335 / 1168)
        blip_fill = picture._inline.xpath(".//pic:blipFill")[0]
        src_rect = OxmlElement("a:srcRect")
        src_rect.set("b", "46314")
        blip_fill.insert(1, src_rect)
    for node in run._r.xpath(".//wp:docPr"):
        node.set("descr", caption_text)
        node.set("title", caption_text.split("。")[0])
    caption = doc.add_paragraph(style=CAPTION)
    add_inline_text(caption, caption_text, size=9)


def style_table(table, rows):
    table.style = "Table Grid"
    table.autofit = True
    compact_sources = bool(rows and rows[0] and rows[0][0] == "来源")
    for r_idx, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if r_idx == 0:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
        for c_idx, cell in enumerate(row.cells):
            if compact_sources:
                set_cell_margins(cell, top=45, start=75, bottom=45, end=75)
            else:
                set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            if r_idx == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), ACCENT)
                tc_pr.append(shd)
            elif r_idx % 2 == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "F5F7FA")
                tc_pr.append(shd)
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = None
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0 if compact_sources else 1.1
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_font(
                        run,
                        size=8.2 if compact_sources else 9,
                        bold=(r_idx == 0),
                        color="FFFFFF" if r_idx == 0 else ACCENT_DARK,
                    )


def add_table(doc: Document, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            paragraph = table.cell(r_idx, c_idx).paragraphs[0]
            paragraph.clear()
            add_inline_text(paragraph, text, size=9.5)
    style_table(table, rows)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)


def add_cover(doc: Document):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("尚硅谷嵌入式项目之 boomPI")
    set_font(run, cn="华文细黑", size=20, bold=True, color=ACCENT_DARK)
    p.paragraph_format.space_after = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("第4章  Linux系统简介")
    set_font(run, cn=FONT_HEADING, size=24, bold=True, color=ACCENT)
    p.paragraph_format.space_after = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("独立审阅稿  v0.1")
    set_font(run, cn=FONT_HEADING, size=13, bold=True, color=GRAY)
    p.paragraph_format.space_after = Pt(48)

    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2026-08-10")
    set_font(run, size=10.5, color=GRAY)
    doc.add_page_break()


def parse_markdown(doc: Document, source: str):
    lines = source.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped == "[[PAGEBREAK]]":
            doc.add_page_break()
            i += 1
            continue
        image = re.fullmatch(r"\[\[IMAGE:(.+?)\|(.+?)\]\]", stripped)
        if image:
            add_figure(doc, image.group(1), image.group(2))
            i += 1
            continue
        if stripped.startswith("| ") and i + 1 < len(lines) and re.match(r"^\|\s*:?-", lines[i + 1].strip()):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-+:?", c.replace(" ", "")) for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                add_table(doc, rows)
            continue
        if stripped.startswith("# "):
            p = doc.add_paragraph(style=H1)
            disable_paragraph_numbering(p)
            heading = stripped[2:]
            add_inline_text(p, heading, size=15)
        elif stripped.startswith("## "):
            p = doc.add_paragraph(style=H2)
            disable_paragraph_numbering(p)
            heading = stripped[3:]
            add_inline_text(p, heading, size=14)
        elif stripped.startswith("> "):
            add_callout(doc, stripped[2:])
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style=BODY)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(-0.45)
            p.paragraph_format.left_indent = Cm(0.75)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run("• ")
            set_font(run, size=10.5, color=ACCENT)
            add_inline_text(p, stripped[2:])
        else:
            p = doc.add_paragraph(style=BODY)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_inline_text(p, stripped)
        i += 1


def scrub_docx_package(path: Path) -> None:
    """Remove template/account residue and media relationships not used by the chapter."""
    with zipfile.ZipFile(path, "r") as source:
        parts = {name: source.read(name) for name in source.namelist()}

    def edit(name: str, callback) -> None:
        if name in parts:
            text = parts[name].decode("utf-8")
            parts[name] = callback(text).encode("utf-8")

    edit(
        "docProps/core.xml",
        lambda text: re.sub(
            r"<cp:lastModifiedBy>.*?</cp:lastModifiedBy>",
            "<cp:lastModifiedBy>尚硅谷嵌入式项目</cp:lastModifiedBy>",
            text,
            flags=re.S,
        ),
    )
    edit("word/settings.xml", lambda text: re.sub(r"<w:attachedTemplate\b[^>]*/>", "", text))
    edit(
        "word/_rels/settings.xml.rels",
        lambda text: re.sub(r"<Relationship\b(?=[^>]*attachedTemplate)[^>]*/>", "", text),
    )

    parts.pop("docProps/custom.xml", None)
    edit(
        "[Content_Types].xml",
        lambda text: re.sub(r"<Override\b(?=[^>]*PartName=\"/docProps/custom\.xml\")[^>]*/>", "", text),
    )
    edit(
        "_rels/.rels",
        lambda text: re.sub(r"<Relationship\b(?=[^>]*Target=\"docProps/custom\.xml\")[^>]*/>", "", text),
    )

    document_xml = parts.get("word/document.xml", b"").decode("utf-8", errors="ignore")
    used_ids = set(re.findall(r'r:(?:id|embed|link)="([^"]+)"', document_xml))
    document_rels = "word/_rels/document.xml.rels"
    if document_rels in parts:
        rel_text = parts[document_rels].decode("utf-8")

        def keep_used_relationship(match) -> str:
            element = match.group(0)
            rel_id = re.search(r'\bId="([^"]+)"', element)
            rel_type = re.search(r'\bType="([^"]+)"', element)
            if rel_id and rel_type and rel_type.group(1).endswith("/image") and rel_id.group(1) not in used_ids:
                return ""
            return element

        rel_text = re.sub(r"<Relationship\b[^>]*/>", keep_used_relationship, rel_text)
        parts[document_rels] = rel_text.encode("utf-8")

    referenced_media = set()
    for name, data in parts.items():
        if not (name.endswith(".rels") and name.startswith("word/")):
            continue
        owner_dir = name.split("/_rels/", 1)[0]
        rel_text = data.decode("utf-8", errors="ignore")
        for target in re.findall(r'\bTarget="([^"]*media/[^"]+)"', rel_text):
            referenced_media.add(normpath(f"{owner_dir}/{target}"))
    for name in list(parts):
        if name.startswith("word/media/") and name not in referenced_media:
            parts.pop(name)

    def scrub_app_properties(text: str) -> str:
        text = re.sub(r"<Template>.*?</Template>", "<Template></Template>", text, flags=re.S)
        text = re.sub(
            r"<Application>.*?</Application>",
            "<Application>Microsoft Office Word</Application>",
            text,
            flags=re.S,
        )
        return re.sub(
            r"<(?:Pages|Words|Characters|CharactersWithSpaces|Lines|Paragraphs)>.*?</(?:Pages|Words|Characters|CharactersWithSpaces|Lines|Paragraphs)>",
            "",
            text,
            flags=re.S,
        )

    edit("docProps/app.xml", scrub_app_properties)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp:
        temp_path = Path(temp.name)
    try:
        with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as output:
            for name, data in parts.items():
                output.writestr(name, data)
        shutil.move(temp_path, path)
    finally:
        temp_path.unlink(missing_ok=True)


def main():
    if not TEMPLATE.exists():
        raise FileNotFoundError(TEMPLATE)
    shutil.copy2(TEMPLATE, OUTPUT)
    doc = Document(OUTPUT)
    clear_document_body(doc)
    configure_styles(doc)
    configure_header_footer(doc)
    doc.core_properties.title = "第4章 Linux系统简介"
    doc.core_properties.subject = "boomPI 独立章节审阅稿"
    doc.core_properties.author = "尚硅谷嵌入式项目"
    doc.core_properties.keywords = "boomPI, Linux, Ubuntu, Buildroot, 目录, 进程"
    add_cover(doc)
    parse_markdown(doc, SOURCE.read_text(encoding="utf-8"))
    doc.save(OUTPUT)
    scrub_docx_package(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
